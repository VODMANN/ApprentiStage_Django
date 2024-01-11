from io import BytesIO
from lesApprentiStage.models import Responsable
from docx import Document as DocxDocument
from .models import *
from django.core.files.base import ContentFile
import os
from django.conf import settings
import logging

from django.template.loader import render_to_string
from django.core.mail import send_mail
from django.http import HttpResponse

from docx2pdf import convert

import os
from django.conf import settings
import subprocess

def convert_docx_to_pdf(input_docx, output_pdf):
    convert(input_docx, output_pdf)

    

logger = logging.getLogger(__name__)


def send_email_with_html_body(subjet: str, receivers: list, template:str, context: dict):
    """ This fonction help to send a customize email to a specific user or set of users."""

    try:
        message = render_to_string(template, context)

        send_mail(
            subjet,
            message,
            settings.EMAIL_HOST_USER,
            receivers,
            fail_silently=False,
            html_message=message
        )

        return True

    except Exception as e:
        print(e)  # Ajoutez ceci pour imprimer l'exception
        logger.error(e)

    return False

def remplacer_texte(doc, recherche, remplacement):
    for p in doc.paragraphs:
        if recherche in p.text:
            for run in p.runs:
                run.text = run.text.replace(recherche, remplacement)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if recherche in run.text:
                            run.text = run.text.replace(recherche, str(remplacement))


def generer_convention(contrat):
    # Charger le modèle de document Word
    doc = DocxDocument(os.path.join(settings.BASE_DIR, 'lesApprentiStage/static/MODELE_CONVENTION_STAGE.docx'))
    # template_path = os.path.join(settings.BASE_DIR, 'lesApprentiStage/static/MODELE_CONVENTION_STAGE.odt')
    # doc = load(template_path)

    # Récupérer le responsable de l'entreprise liée au contrat
    responsable = Responsable.objects.filter(entreprise=contrat.entreprise).first()

    # Définir le contexte de remplacement
    context = {
        '{{nom_etudiant}}': contrat.etudiant.nomEtu,
        '{{prenom_etudiant}}': contrat.etudiant.prenomEtu,
        '{{nom_entreprise}}': contrat.entreprise.nomEnt,
        '{{adresse_entreprise}}': contrat.entreprise.adresseEnt,
        '{{nom_responsable}}': responsable.nomResp if responsable else '',
        '{{email_responsable}}': responsable.emailResp if responsable else '',
        '{{tel_responsable}}': responsable.telResp if responsable else '',
        '{{adresse_etudiant}}': contrat.etudiant.adresseEtu,
        '{{nom_tuteur}}': contrat.tuteur.nomTuteur,
        '{{prenom_tuteur}}': contrat.tuteur.prenomTuteur,
        '{{metier_tuteur}}': contrat.tuteur.metierTuteur,
        '{{mail_tuteur}}': contrat.tuteur.emailTuteur,
        '{{tel_tuteur}}': contrat.tuteur.telTuteur,
        '{{sexeM_etu}}': 'M ☑ ' if contrat.etudiant.civiliteEtu == 'M' else 'M ◻ ',
        '{{sexeF_etu}}': 'F ◻ ' if contrat.etudiant.civiliteEtu == 'M' else 'F ☑ ',
        '{{email_etudiant}}': contrat.etudiant.mailEtu,
        '{{tel_etudiant}}': contrat.etudiant.telEtu,
        '{{dateNaissance_etudiant}}': contrat.etudiant.dateNEtu,
        '{{titre_contrat}}': contrat.titre,
        '{{promo}}':contrat.etudiant.promo.nomPromo,
    }

    # Remplacer le texte dans les paragraphes et les tables
    for key, value in context.items():
        remplacer_texte(doc, key, value)

    # Sauvegarder le document dans un buffer en mémoire
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # Créer un ContentFile à partir du buffer
    file_name = f'Convention_{contrat.etudiant.nomEtu}_{contrat.etudiant.prenomEtu}.docx'
    django_file = ContentFile(file_stream.read(), name=file_name)
    
    # Créer et enregistrer le nouveau document dans le modèle Document
    nouveau_document = Document(titre=file_name, fichier=django_file, contrat=contrat)
    nouveau_document.save()

    return nouveau_document.fichier.path



def insert(request):
    # Création des instances Utilisateur (Utilisateur personnalisé)
    etudiant = Utilisateur.objects.create_user(
        username='etudiant123',
        email='etudiant@example.com',
        password='etudiantpassword',
        first_name='John',
        last_name='Doe',
        type_utilisateur='etudiant'  # Ajoutez le champ type_utilisateur selon votre modèle
    )

    enseignant = Utilisateur.objects.create_user(
        username='enseignant456',
        email='enseignant@example.com',
        password='enseignantpassword',
        first_name='Alice',
        last_name='Smith',
        type_utilisateur='enseignant'  # Ajoutez le champ type_utilisateur selon votre modèle
    )

    secretaire = Utilisateur.objects.create_user(
        username='secretaire789',
        email='secretaire@example.com',
        password='secretairepassword',
        first_name='Bob',
        last_name='Johnson',
        type_utilisateur='secretaire'  # Ajoutez le champ type_utilisateur selon votre modèle
    )
    
    profil_enseignant = ProfilEnseignant.objects.create(
        utilisateur=enseignant,
        numHarpege='123ABC',
        nomEnseignant='Smith',
        prenomEnseignant='Alice',
        mailEnseignant='enseignant@example.com',
        roleEnseignant='chef_departement'
    )
    
    # Création des instances Departement, Entreprise, Theme
    departement_info = Departement.objects.create(
        nomDep='Département Informatique',
        adresseDep='123 Rue des Développeurs',
        chef=profil_enseignant  # Vous pouvez attribuer un enseignant existant comme chef de département
    )
    
    # Création des instances Promo
    promo_info = Promo.objects.create(
        nomPromo='1B',
        anneeScolaire='2023-2024',
        departement=departement_info
    )
    
    # Création des instances ProfilEtudiant, ProfilEnseignant, ProfilSecretaire
    profil_etudiant = ProfilEtudiant.objects.create(
        utilisateur=etudiant,
        nomEtu='Doe',
        prenomEtu='John',
        numEtu='E12345',
        civiliteEtu='M.',
        adresseEtu='789 Avenue des Étudiants',
        cpEtu=54321,
        villeEtu='Ville Étudiant',
        telEtu='0123456789',
        promo=promo_info,
        idDepartement=departement_info
    )



    profil_secretaire = ProfilSecretaire.objects.create(
        utilisateur=secretaire,
        numSec='S123'
    )




    entreprise_info = Entreprise.objects.create(
        numSiret='123456789',
        nomEnt='Entreprise XYZ',
        adresseEnt='456 Rue des Entreprises',
        cpEnt=12345,
        villeEnt='Ville Entreprise'
    )

    theme_info = Theme.objects.create(
        nomTheme='Développement Web'
    )



    
    
        # Création des instances EnseignantPromo
    enseignant_promo = EnseignantPromo.objects.create(
        enseignant=profil_enseignant,
        promo=promo_info
    )

    # Création des instances Responsable, Tuteur, Contrat
    responsable_entreprise = Responsable.objects.create(
        nomResp='Nom Responsable',
        prenomResp='Prénom Responsable',
        emailResp='responsable@entreprise.com',
        entreprise=entreprise_info
    )

    tuteur_entreprise = Tuteur.objects.create(
        nomTuteur='Nom Tuteur',
        prenomTuteur='Prénom Tuteur',
        metierTuteur='Métier Tuteur',
        telTuteur='9876543210',
        emailTuteur='tuteur@entreprise.com',
        entreprise=entreprise_info
    )

    contrat_info = Contrat.objects.create(
        type='Type Contrat',
        description='Description Contrat',
        etat='État Contrat',
        gratification='Gratification Contrat',
        dateDeb=timezone.now(),
        dateFin=timezone.now() + timezone.timedelta(days=90),  # Exemple : Contrat pour 90 jours
        estValide=True,
        etudiant=profil_etudiant,
        enseignant=profil_enseignant,
        tuteur=tuteur_entreprise,
        theme=theme_info,
        entreprise=entreprise_info,
        enFrance=True
    )


    
    # Création des instances Offre, Salle, Soutenance, Document, Evaluation
    offre_info = Offre.objects.create(
        titre='Titre Offre',
        description='Description Offre',
        mailRh='rh@entreprise.com',
        competences='Compétences requises',
        duree='3 mois',
        datePublication=timezone.now(),
        entreprise=entreprise_info,
        theme=theme_info,
        estPublie=False
    )

    salle_info = Salle.objects.create(
        numero='Salle A'
    )

    soutenance_info = Soutenance.objects.create(
        dateSoutenance=timezone.now(),
        heureSoutenance=timezone.now().time(),
        salle=salle_info,
        idContrat=contrat_info,
        candide=profil_enseignant,
        estDistanciel=False
    )

    document_info = Document.objects.create(
        titre='Titre Document',
        fichier='chemin/vers/fichier.pdf',
        contrat=contrat_info
    )

    evaluation_info = Evaluation.objects.create(
        contrat=contrat_info,
        enseignant=profil_enseignant,
        note=4.5,
        commentaire='Commentaire évaluation'
    )
    
    return HttpResponse("Données insérées avec succès !")
